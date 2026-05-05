"""
测试文档生成脚本
生成 3 个 PDF 文件和 3 个 DOCX 文件，用于知识库上传测试
运行方式：python generate_docs.py
"""
import os
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 输出目录为当前脚本所在目录
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ========================================
# PDF 文档内容定义
# ========================================

PDF_DOCS = {
    "公司薪酬福利制度": [
        ("公司薪酬福利制度", True),
        ("", False),
        ("第一章 总则", True),
        ("", False),
        ("第一条 为建立科学合理的薪酬福利体系，充分调动员工的工作积极性和创造性，吸引和保留优秀人才，根据国家相关法律法规，结合公司实际情况，特制定本制度。", False),
        ("", False),
        ("第二条 本制度适用于公司全体正式员工。试用期员工的薪酬按照录用通知书中约定的标准执行。", False),
        ("", False),
        ("第二章 薪酬结构", True),
        ("", False),
        ("第三条 员工薪酬由以下部分组成：", False),
        ("（一）基本工资：根据岗位等级确定，是员工薪酬的基础部分。", False),
        ("（二）岗位津贴：根据岗位职责和工作难度确定的补贴。", False),
        ("（三）绩效奖金：根据员工个人绩效考核结果发放，体现多劳多得。", False),
        ("（四）全勤奖金：当月无迟到、早退、旷工、请假记录的，发放全勤奖金500元/月。", False),
        ("（五）年终奖金：根据公司年度经营业绩和个人考核结果发放。", False),
        ("", False),
        ("第四条 岗位等级与薪资范围：", False),
        ("P1-P3（初级）：基本工资 6000-10000元/月", False),
        ("P4-P6（中级）：基本工资 10000-18000元/月", False),
        ("P7-P9（高级）：基本工资 18000-30000元/月", False),
        ("P10及以上（专家级）：基本工资 30000元以上/月", False),
        ("", False),
        ("第三章 薪酬发放", True),
        ("", False),
        ("第五条 工资发放时间为每月15日，如遇节假日则提前至最近的工作日发放。", False),
        ("", False),
        ("第六条 工资通过银行转账方式发放至员工指定的个人银行账户。", False),
        ("", False),
        ("第七条 公司按照国家规定代扣代缴个人所得税和社会保险个人缴纳部分。", False),
        ("", False),
        ("第四章 福利制度", True),
        ("", False),
        ("第八条 社会保险：公司为员工缴纳养老保险、医疗保险、失业保险、工伤保险和生育保险。", False),
        ("", False),
        ("第九条 住房公积金：公司按照员工工资的12%缴纳住房公积金。", False),
        ("", False),
        ("第十条 补充福利：", False),
        ("（一）餐饮补贴：每月500元餐补，随工资一同发放。", False),
        ("（二）交通补贴：每月300元交通补贴。", False),
        ("（三）通讯补贴：根据岗位需要，每月100-500元。", False),
        ("（四）节日福利：春节、中秋节、端午节等传统节日发放节日礼品或礼金。", False),
        ("（五）生日福利：员工生日当月发放200元生日礼金和生日蛋糕。", False),
        ("（六）体检福利：每年组织一次免费全面体检。", False),
        ("（七）培训福利：公司提供内部培训和外部培训机会，费用由公司承担。", False),
        ("（八）团建活动：每季度组织一次团队建设活动。", False),
        ("", False),
        ("第五章 薪酬调整", True),
        ("", False),
        ("第十一条 公司每年4月进行一次年度薪酬调整，调整依据包括：", False),
        ("（一）年度绩效考核结果", False),
        ("（二）市场薪酬水平变化", False),
        ("（三）公司经营业绩状况", False),
        ("（四）岗位变动情况", False),
        ("", False),
        ("第十二条 员工晋升时，薪酬按照新岗位等级标准进行调整。", False),
        ("", False),
        ("第六章 附则", True),
        ("", False),
        ("第十三条 本制度由人力资源部负责解释和修订。", False),
        ("第十四条 本制度自2024年1月1日起执行。", False),
    ],
    "数据库设计规范": [
        ("数据库设计规范", True),
        ("", False),
        ("一、概述", True),
        ("", False),
        ("本规范旨在统一公司内部数据库设计标准，提高数据库性能和可维护性。适用于MySQL数据库的设计与开发。", False),
        ("", False),
        ("二、命名规范", True),
        ("", False),
        ("2.1 数据库命名", False),
        ("数据库名以db_开头，使用小写字母和下划线，如：db_enterprise_qa、db_user_center。", False),
        ("", False),
        ("2.2 表命名", False),
        ("表名以t_开头，使用小写字母和下划线，如：t_user、t_order_detail。", False),
        ("关联表命名格式：t_关联表1_关联表2，如：t_user_role。", False),
        ("", False),
        ("2.3 字段命名", False),
        ("字段名使用小写字母和下划线，如：user_name、create_time。", False),
        ("主键统一命名为id，使用自增整型。", False),
        ("外键命名格式：关联表名_id，如：user_id、order_id。", False),
        ("时间字段统一以_time结尾，如：create_time、update_time。", False),
        ("状态字段统一命名为status，使用TINYINT类型。", False),
        ("", False),
        ("2.4 索引命名", False),
        ("普通索引：idx_表名_字段名，如：idx_user_username。", False),
        ("唯一索引：uk_表名_字段名，如：uk_user_email。", False),
        ("联合索引：idx_表名_字段1_字段2，如：idx_order_user_id_status。", False),
        ("", False),
        ("三、表设计规范", True),
        ("", False),
        ("3.1 基本要求", False),
        ("所有表必须使用InnoDB存储引擎。", False),
        ("字符集统一使用utf8mb4，排序规则使用utf8mb4_unicode_ci。", False),
        ("所有表必须有主键，推荐使用自增整型主键。", False),
        ("所有表必须包含create_time和update_time字段。", False),
        ("表和字段必须添加COMMENT注释。", False),
        ("", False),
        ("3.2 字段类型选择", False),
        ("整数类型：优先使用INT，大范围使用BIGINT。", False),
        ("小数类型：金额使用DECIMAL(10,2)，禁止使用FLOAT和DOUBLE。", False),
        ("字符串类型：固定长度使用CHAR，可变长度使用VARCHAR，大文本使用TEXT。", False),
        ("时间类型：统一使用DATETIME类型。", False),
        ("布尔类型：使用TINYINT(1)，0表示否，1表示是。", False),
        ("", False),
        ("3.3 字段约束", False),
        ("所有字段尽量设置NOT NULL，用默认值替代NULL。", False),
        ("VARCHAR类型字段默认值设为空字符串。", False),
        ("数值类型字段默认值设为0。", False),
        ("", False),
        ("四、索引规范", True),
        ("", False),
        ("4.1 索引原则", False),
        ("WHERE条件、JOIN关联、ORDER BY排序中频繁使用的字段应建立索引。", False),
        ("区分度低的字段（如性别、状态）不宜单独建立索引。", False),
        ("单表索引数量不超过5个。", False),
        ("联合索引遵循最左前缀原则。", False),
        ("", False),
        ("4.2 禁止事项", False),
        ("禁止在索引列上使用函数或表达式。", False),
        ("禁止使用SELECT *，只查询需要的字段。", False),
        ("禁止在WHERE条件中对索引列使用!=或NOT IN。", False),
        ("", False),
        ("五、SQL编写规范", True),
        ("", False),
        ("5.1 查询规范", False),
        ("使用参数化查询，禁止SQL拼接。", False),
        ("大批量数据操作使用分批处理。", False),
        ("避免使用子查询，优先使用JOIN。", False),
        ("分页查询必须带有ORDER BY。", False),
        ("", False),
        ("5.2 事务规范", False),
        ("事务要尽可能短小。", False),
        ("事务中避免进行外部服务调用。", False),
        ("必须处理事务的异常回滚。", False),
        ("", False),
        ("六、安全规范", True),
        ("", False),
        ("数据库账号权限遵循最小权限原则。", False),
        ("生产环境禁止使用root账号连接。", False),
        ("敏感数据（如密码）必须加密存储。", False),
        ("定期进行数据库备份，备份文件加密保存。", False),
    ],
    "企业邮箱配置说明": [
        ("企业邮箱配置说明", True),
        ("", False),
        ("一、概述", True),
        ("", False),
        ("公司统一使用企业邮箱作为内部沟通和对外商务联系的官方邮箱。本文档介绍企业邮箱在各类设备和客户端上的配置方法。", False),
        ("", False),
        ("邮箱格式：姓名拼音@company.com（如：zhangsan@company.com）", False),
        ("邮箱初始密码：入职时由IT部门设置并告知", False),
        ("邮箱管理后台：https://mail.company.com/admin", False),
        ("", False),
        ("二、网页版使用", True),
        ("", False),
        ("2.1 登录方式", False),
        ("访问 https://mail.company.com，输入完整邮箱地址和密码即可登录。", False),
        ("", False),
        ("2.2 功能说明", False),
        ("网页版支持收发邮件、通讯录管理、日历管理等全部功能。", False),
        ("支持邮件搜索、文件夹管理、邮件标签等高级功能。", False),
        ("单封邮件附件大小限制为50MB。", False),
        ("邮箱容量为10GB，超过80%使用率时系统会发出提醒。", False),
        ("", False),
        ("三、Outlook客户端配置", True),
        ("", False),
        ("3.1 自动配置", False),
        ('打开Outlook，选择"文件 - 添加账户"。', False),
        ('输入邮箱地址，点击"连接"。', False),
        ("系统将自动发现服务器设置并完成配置。", False),
        ("", False),
        ("3.2 手动配置", False),
        ("如自动配置失败，请使用以下参数手动配置：", False),
        ("", False),
        ("IMAP接收服务器：", False),
        ("  服务器地址：imap.company.com", False),
        ("  端口：993", False),
        ("  加密方式：SSL/TLS", False),
        ("", False),
        ("SMTP发送服务器：", False),
        ("  服务器地址：smtp.company.com", False),
        ("  端口：465", False),
        ("  加密方式：SSL/TLS", False),
        ("  需要身份验证：是", False),
        ("", False),
        ("四、手机客户端配置", True),
        ("", False),
        ("4.1 iOS设备", False),
        ('打开"设置 - 邮件 - 账户 - 添加账户"。', False),
        ('选择"其他"，然后选择"添加邮件账户"。', False),
        ("输入姓名、邮箱地址和密码。", False),
        ("选择IMAP类型，填写服务器信息（同上述Outlook手动配置参数）。", False),
        ("", False),
        ("4.2 Android设备", False),
        ("打开设备自带的邮件应用或Gmail应用。", False),
        ('选择"添加账户 - 其他"。', False),
        ("输入邮箱地址和密码。", False),
        ("选择IMAP类型，填写服务器信息。", False),
        ("", False),
        ("五、常见问题", True),
        ("", False),
        ("Q1: 无法收到邮件怎么办？", False),
        ("A1: 请检查网络连接，确认服务器设置是否正确，检查垃圾邮件文件夹。", False),
        ("", False),
        ("Q2: 发送邮件提示认证失败？", False),
        ("A2: 请确认SMTP服务器已开启身份验证，用户名填写完整邮箱地址。", False),
        ("", False),
        ("Q3: 如何修改邮箱密码？", False),
        ('A3: 登录网页版邮箱，在"设置 - 安全设置"中修改密码。密码须包含大小写字母、数字，长度不少于8位。', False),
        ("", False),
        ("Q4: 邮箱容量满了怎么办？", False),
        ("A4: 请清理不需要的邮件和附件，或将重要邮件归档到本地。如需扩容请联系IT部门申请。", False),
        ("", False),
        ("六、技术支持", True),
        ("", False),
        ("如遇到无法解决的邮箱问题，请联系IT技术支持：", False),
        ("电话：内线8888", False),
        ("邮箱：it-support@company.com", False),
        ("工单系统：http://helpdesk.company.com", False),
    ],
}


# ========================================
# DOCX 文档内容定义
# ========================================

DOCX_DOCS = {
    "员工行为规范手册": {
        "sections": [
            {
                "title": "员工行为规范手册",
                "level": 0,
                "content": []
            },
            {
                "title": "第一章 总则",
                "level": 1,
                "content": [
                    "第一条 为规范员工行为，维护公司良好的工作秩序和企业形象，营造和谐的工作环境，特制定本手册。",
                    "第二条 本手册适用于公司全体员工，包括正式员工、试用期员工和实习生。",
                    "第三条 全体员工须认真学习并严格遵守本手册的各项规定。",
                ]
            },
            {
                "title": "第二章 职业道德",
                "level": 1,
                "content": [
                    "第四条 基本要求：",
                    "（一）爱岗敬业，忠于职守，勤勉工作；",
                    "（二）诚实守信，正直公正，廉洁自律；",
                    "（三）团结协作，相互尊重，和谐共处；",
                    "（四）积极进取，勇于创新，追求卓越。",
                    "",
                    "第五条 保密义务：",
                    "（一）员工须对公司的商业秘密、技术秘密和客户信息严格保密；",
                    "（二）未经授权，不得对外披露公司内部文件、数据和决策信息；",
                    "（三）离职后仍须遵守保密协议中约定的保密义务。",
                    "",
                    "第六条 利益冲突：",
                    "（一）员工不得利用职务之便谋取个人私利；",
                    "（二）不得在与公司有业务往来的单位兼职或持有股份；",
                    "（三）如存在潜在利益冲突，须及时向直属领导和人力资源部报告。",
                ]
            },
            {
                "title": "第三章 工作纪律",
                "level": 1,
                "content": [
                    "第七条 考勤纪律：",
                    "（一）按时上下班，不迟到、不早退、不旷工；",
                    "（二）请假须提前申请并获得批准；",
                    "（三）外出办公须在OA系统中登记。",
                    "",
                    "第八条 办公纪律：",
                    "（一）工作时间内专注工作，不做与工作无关的事务；",
                    "（二）保持工作区域整洁有序；",
                    "（三）公共区域轻声交流，不影响他人工作；",
                    "（四）会议期间将手机调至静音模式。",
                    "",
                    "第九条 信息安全：",
                    "（一）妥善保管个人账号和密码，不得与他人共用；",
                    "（二）公司电脑不得安装未授权的软件；",
                    "（三）重要文件须定期备份；",
                    "（四）离开工位时须锁定电脑屏幕。",
                ]
            },
            {
                "title": "第四章 着装与礼仪",
                "level": 1,
                "content": [
                    "第十条 着装要求：",
                    "（一）日常办公着装整洁得体，符合职业形象；",
                    "（二）接待客户或参加正式会议时，须着正装；",
                    "（三）不得穿拖鞋、背心等过于随意的服装进入办公区域。",
                    "",
                    "第十一条 商务礼仪：",
                    "（一）接待来访客户须热情友好，主动问好并引导；",
                    "（二）商务交往中注意言行举止，维护公司形象；",
                    "（三）商务用餐遵循基本用餐礼仪。",
                ]
            },
            {
                "title": "第五章 奖惩制度",
                "level": 1,
                "content": [
                    "第十二条 奖励类型：",
                    "（一）通报表扬：对工作表现突出的员工给予全公司通报表扬；",
                    "（二）奖金奖励：对做出重大贡献的员工给予一次性奖金奖励；",
                    "（三）晋升奖励：对持续表现优异的员工优先考虑晋升。",
                    "",
                    "第十三条 处罚类型：",
                    "（一）口头警告：适用于轻微违规行为；",
                    "（二）书面警告：适用于较严重或重复违规行为；",
                    "（三）降职降薪：适用于严重违规行为；",
                    "（四）解除合同：适用于违反法律法规或造成重大损失的行为。",
                ]
            },
            {
                "title": "第六章 附则",
                "level": 1,
                "content": [
                    "第十四条 本手册由人力资源部负责解释和修订。",
                    "第十五条 本手册自发布之日起执行，原有相关规定与本手册不一致的，以本手册为准。",
                ]
            },
        ]
    },
    "Git版本管理规范": {
        "sections": [
            {
                "title": "Git版本管理规范",
                "level": 0,
                "content": []
            },
            {
                "title": "一、概述",
                "level": 1,
                "content": [
                    "本规范旨在统一公司内部Git版本管理流程，提高团队协作效率，保证代码质量和版本管理的规范性。所有开发人员须严格遵循本规范。",
                ]
            },
            {
                "title": "二、分支管理策略",
                "level": 1,
                "content": [
                    "2.1 主要分支：",
                    "main分支：生产环境代码，仅通过合并请求更新，禁止直接提交。",
                    "develop分支：开发主分支，集成各功能分支的代码。",
                    "",
                    "2.2 辅助分支：",
                    "feature/xxx分支：功能开发分支，从develop分支创建，开发完成后合并回develop。",
                    "hotfix/xxx分支：紧急修复分支，从main分支创建，修复完成后同时合并回main和develop。",
                    "release/x.x.x分支：版本发布分支，从develop分支创建，用于发布前的测试和修复。",
                    "",
                    "2.3 分支命名规范：",
                    "feature/user-login：用户登录功能",
                    "feature/order-export：订单导出功能",
                    "hotfix/fix-login-bug：修复登录Bug",
                    "release/1.2.0：1.2.0版本发布",
                ]
            },
            {
                "title": "三、提交规范",
                "level": 1,
                "content": [
                    "3.1 提交信息格式：",
                    "type(scope): description",
                    "",
                    "3.2 type类型说明：",
                    "feat：新功能",
                    "fix：Bug修复",
                    "docs：文档更新",
                    "style：代码格式调整（不影响功能）",
                    "refactor：代码重构（不是新功能也不是Bug修复）",
                    "test：添加或修改测试",
                    "chore：构建工具或辅助工具的变动",
                    "",
                    "3.3 提交示例：",
                    "feat(user): 添加用户登录功能",
                    "fix(order): 修复订单金额计算错误",
                    "docs(readme): 更新项目部署文档",
                    "",
                    "3.4 提交原则：",
                    "每次提交只做一件事情，保持提交的原子性。",
                    "提交信息要准确描述变更内容。",
                    "不要提交未完成的代码到共享分支。",
                    "提交前须运行代码检查和单元测试。",
                ]
            },
            {
                "title": "四、代码合并流程",
                "level": 1,
                "content": [
                    "4.1 合并请求（Merge Request）：",
                    "所有代码合并必须通过合并请求进行。",
                    "合并请求须至少1名其他开发人员进行代码审查。",
                    "审查通过后方可合并。",
                    "",
                    "4.2 代码审查要点：",
                    "代码逻辑是否正确。",
                    "是否遵循编码规范。",
                    "是否有潜在的性能问题。",
                    "是否有安全漏洞。",
                    "是否有足够的注释和文档。",
                    "",
                    "4.3 冲突解决：",
                    "合并前须先将目标分支的最新代码合并到当前分支。",
                    "冲突须由功能开发者自行解决。",
                    "解决冲突后须重新测试。",
                ]
            },
            {
                "title": "五、版本发布流程",
                "level": 1,
                "content": [
                    "5.1 版本号规范（语义化版本）：",
                    "主版本号.次版本号.修订号（如1.2.3）",
                    "主版本号：不兼容的API变更",
                    "次版本号：向下兼容的功能新增",
                    "修订号：向下兼容的问题修复",
                    "",
                    "5.2 发布流程：",
                    "从develop分支创建release分支。",
                    "在release分支上进行测试和Bug修复。",
                    "测试通过后合并到main分支并打Tag。",
                    "同时将release分支合并回develop分支。",
                    "删除release分支。",
                ]
            },
            {
                "title": "六、其他规定",
                "level": 1,
                "content": [
                    "禁止将敏感信息（密码、密钥、Token等）提交到代码仓库。",
                    "使用.gitignore忽略编译产物、依赖目录和IDE配置文件。",
                    "大文件（超过10MB）不要提交到Git仓库，使用Git LFS或其他存储方案。",
                    "定期清理已合并的分支，保持仓库整洁。",
                ]
            },
        ]
    },
    "新员工入职指南": {
        "sections": [
            {
                "title": "新员工入职指南",
                "level": 0,
                "content": []
            },
            {
                "title": "一、欢迎加入",
                "level": 1,
                "content": [
                    "欢迎您加入公司大家庭！本指南将帮助您快速了解公司环境、办理入职手续、熟悉工作流程，顺利开启您的职业新旅程。",
                ]
            },
            {
                "title": "二、入职手续办理",
                "level": 1,
                "content": [
                    "2.1 入职当天须携带的材料：",
                    "（一）身份证原件及复印件2份",
                    "（二）学历证书原件及复印件1份",
                    "（三）学位证书原件及复印件1份",
                    "（四）前公司离职证明原件",
                    "（五）近期一寸免冠照片2张",
                    "（六）银行卡复印件1份（用于工资发放）",
                    "（七）体检报告（三个月内有效）",
                    "",
                    "2.2 入职办理流程：",
                    "第一步：到人力资源部报到，提交入职材料；",
                    "第二步：签订劳动合同和保密协议；",
                    "第三步：领取工牌、门禁卡、办公用品；",
                    "第四步：IT部门开通电脑账号、邮箱、OA系统等；",
                    "第五步：由部门负责人或指定导师带领参观办公环境；",
                    "第六步：参加新员工入职培训。",
                ]
            },
            {
                "title": "三、办公环境介绍",
                "level": 1,
                "content": [
                    "3.1 办公区域：",
                    "公司位于科技园A座，共占用5-8层。",
                    "5楼：行政部、人力资源部、财务部",
                    "6楼：产品部、设计部",
                    "7楼：技术研发部",
                    "8楼：市场部、销售部、会议室",
                    "",
                    "3.2 公共设施：",
                    "餐厅：位于1楼，提供早餐和午餐，营业时间7:30-9:00和11:30-13:00。",
                    "健身房：位于B座1楼，7:00-22:00开放，员工免费使用。",
                    "休息区：每层楼设有茶水间和休息区，提供免费咖啡、茶和零食。",
                    "会议室：使用前须通过OA系统预约。",
                    "停车场：地下停车场B1-B2层，新员工可向行政部申请停车位。",
                ]
            },
            {
                "title": "四、IT系统开通",
                "level": 1,
                "content": [
                    "入职后IT部门将为您开通以下系统账号：",
                    "",
                    "4.1 电脑登录账号",
                    "用户名：工号",
                    "初始密码：身份证后6位",
                    "首次登录须修改密码",
                    "",
                    "4.2 企业邮箱",
                    "格式：姓名拼音@company.com",
                    "网页登录：https://mail.company.com",
                    "",
                    "4.3 OA办公系统",
                    "地址：http://oa.company.com",
                    "用于考勤打卡、请假申请、流程审批等",
                    "",
                    "4.4 项目管理平台",
                    "地址：http://pm.company.com",
                    "用于项目任务管理和协作",
                    "",
                    "4.5 代码仓库（技术岗位）",
                    "地址：https://git.company.com",
                    "用于代码版本管理",
                ]
            },
            {
                "title": "五、试用期说明",
                "level": 1,
                "content": [
                    "5.1 试用期时长：",
                    "劳动合同期限1年以上不满3年的，试用期为2个月。",
                    "劳动合同期限3年以上的，试用期为3个月。",
                    "",
                    "5.2 试用期考核：",
                    "试用期内由直属主管和导师共同对新员工进行考核。",
                    "考核内容包括：工作能力、学习能力、团队协作、职业素养等。",
                    "试用期结束前一周，由人力资源部组织转正评审。",
                    "",
                    "5.3 试用期薪资：",
                    "试用期工资为正式工资的80%。",
                    "试用期内享有法定节假日、社保等基本权益。",
                ]
            },
            {
                "title": "六、常见问题",
                "level": 1,
                "content": [
                    "Q1: 忘记带工牌怎么办？",
                    "A1: 到前台登记身份信息，领取临时访客卡。",
                    "",
                    "Q2: 电脑账号被锁定怎么办？",
                    "A2: 联系IT部门（内线8888）进行解锁。",
                    "",
                    "Q3: 如何申请办公用品？",
                    "A3: 在OA系统中提交办公用品申请，由行政部统一采购配发。",
                    "",
                    "Q4: 公司有哪些培训机会？",
                    "A4: 公司提供新人培训、专业技能培训、管理能力培训等，详情可咨询人力资源部。",
                ]
            },
        ]
    },
}


def create_pdf_docs():
    """使用fpdf2生成PDF测试文档"""
    font_path = "C:/Windows/Fonts/simhei.ttf"

    for filename, content_lines in PDF_DOCS.items():
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)

        if os.path.exists(font_path):
            pdf.add_font("simhei", "", font_path)
            font_name = "simhei"
        else:
            font_name = "Helvetica"

        for text, is_title in content_lines:
            if not text:
                pdf.ln(3)
                continue
            if is_title:
                pdf.set_font(font_name, "", 14)
                pdf.cell(w=0, h=10, text=text, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
            else:
                pdf.set_font(font_name, "", 10)
                pdf.multi_cell(w=0, h=6, text=text, new_x="LMARGIN", new_y="NEXT")

        output_path = os.path.join(OUTPUT_DIR, f"{filename}.pdf")
        pdf.output(output_path)
        print(f"已生成: {output_path}")


def create_docx_docs():
    """使用python-docx生成DOCX测试文档"""
    for filename, doc_data in DOCX_DOCS.items():
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        for section in doc_data["sections"]:
            title = section["title"]
            level = section["level"]

            if level == 0:
                # 文档主标题
                heading = doc.add_heading(title, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # 章节标题
                doc.add_heading(title, level=1)

            # 添加内容段落
            for line in section["content"]:
                if line == "":
                    doc.add_paragraph("")
                else:
                    p = doc.add_paragraph(line)
                    p.style.font.size = Pt(11)

        output_path = os.path.join(OUTPUT_DIR, f"{filename}.docx")
        doc.save(output_path)
        print(f"已生成: {output_path}")


if __name__ == '__main__':
    print("开始生成测试文档...\n")
    print("=== 生成PDF文档 ===")
    create_pdf_docs()
    print("\n=== 生成DOCX文档 ===")
    create_docx_docs()
    print(f"\n全部生成完成！文件位于: {OUTPUT_DIR}")
